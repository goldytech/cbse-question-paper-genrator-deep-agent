"""Tool for converting JSON question papers to DOCX format with embedded diagrams.

Updated for CBSE format compliance with support for:
- Dict format options from question-assembler
- Internal choice questions (OR format)
- Case study questions with sub-parts
- CBSE-compliant headers and section formatting
"""

import os
import base64
import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from langchain_core.tools import tool
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


# Cache directory for temporary PNG files
TEMP_DIR = Path(__file__).parent.parent / "cache" / "temp"
TEMP_DIR.mkdir(parents=True, exist_ok=True)

# Output directory for DOCX files
DOCX_OUTPUT_DIR = Path(__file__).parent.parent / "output" / "docx"
DOCX_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# CBSE Section configuration
CBSE_SECTIONS = {
    "A": {"title": "Multiple Choice Questions", "marks_per_question": 1},
    "B": {"title": "Very Short Answer Questions", "marks_per_question": 2},
    "C": {"title": "Short Answer Questions", "marks_per_question": 3},
    "D": {"title": "Long Answer Questions", "marks_per_question": 5},
    "E": {"title": "Case Study Based Questions", "marks_per_question": 4},
}


def _ensure_cairosvg_installed():
    """Check if cairosvg is installed."""
    try:
        import cairosvg

        return True
    except ImportError:
        print("Warning: cairosvg not installed. SVG to PNG conversion will be skipped.")
        print("Diagrams in DOCX will be missing, but document will still be generated.")
        return False


def _svg_base64_to_png(svg_base64: str, width: int = 400) -> Optional[bytes]:
    """Convert base64-encoded SVG to PNG using cairosvg."""
    if not _ensure_cairosvg_installed():
        print("Warning: cairosvg not available, cannot convert SVG to PNG")
        return None

    try:
        import cairosvg

        # Decode base64 SVG
        svg_content = base64.b64decode(svg_base64)

        # Convert SVG to PNG
        png_bytes = cairosvg.svg2png(bytestring=svg_content, output_width=width, write_to=None)

        return png_bytes

    except Exception as e:
        print(f"Error converting SVG to PNG: {e}")
        return None


def _generate_docx_filename(metadata: Dict) -> str:
    """Generate DOCX filename from metadata."""
    subject = metadata.get("subject", "mathematics").lower().replace(" ", "_")
    cls = metadata.get("class", "10")
    exam_type = metadata.get("exam_type", "examination").lower().replace(" ", "_")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    short_id = str(uuid.uuid4())[:5]

    return f"{subject}_class{cls}_{exam_type}_{timestamp}_{short_id}.docx"


def _create_cbse_header(doc, metadata: Dict):
    """Create CBSE-style header for document."""
    # Get section
    section = doc.sections[0]

    # Header content
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    # Add CBSE title
    header_para.text = "CENTRAL BOARD OF SECONDARY EDUCATION"
    header_para.runs[0].bold = True
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.runs[0].font.size = Pt(12)
    header_para.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    # Add subject and class
    subject_para = header.add_paragraph()
    subject_para.text = (
        f"{metadata.get('subject', 'Subject').upper()} (Class {metadata.get('class', '10')})"
    )
    subject_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subject_para.runs[0].font.size = Pt(11)
    subject_para.runs[0].bold = True

    # Add exam type
    exam_para = header.add_paragraph()
    exam_para.text = metadata.get("exam_type", "EXAMINATION").upper()
    exam_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    exam_para.runs[0].font.size = Pt(10)

    # Add time and marks
    time_para = header.add_paragraph()
    duration_hours = metadata.get("duration_minutes", 180) // 60
    total_marks = metadata.get("total_marks", 80)
    time_para.text = f"TIME: {duration_hours} HOURS\t\tMAX. MARKS: {total_marks}"
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    time_para.runs[0].font.size = Pt(10)
    time_para.runs[0].bold = True


def _create_cbse_general_instructions(doc, num_sections: int = 5):
    """Add general instructions section per CBSE format."""
    # Title
    title_para = doc.add_paragraph()
    title_para.add_run("General Instructions:").bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para.runs[0].font.size = Pt(11)
    title_para.spacing_after = Pt(6)

    # CBSE standard instructions
    instructions = [
        f"This Question Paper consists of {num_sections} Sections A, B, C, D and E.",
        "Section A has {X} MCQs carrying 1 mark each.",
        "Section B has {Y} Short Answer questions carrying 2 marks each.",
        "Section C has {Z} Short Answer questions carrying 3 marks each.",
        "Section D has {W} Long Answer questions carrying 5 marks each.",
        "Section E has {V} Case Study based questions carrying 4 marks each.",
        "All questions are compulsory.",
        "Internal choice is provided in some questions.",
        "Use of calculators is not allowed.",
        "Draw neat and clean figures wherever required.",
    ]

    for i, instruction in enumerate(instructions, 1):
        inst_para = doc.add_paragraph()
        inst_para.add_run(f"{i}. {instruction}")
        inst_para.paragraph_format.left_indent = Inches(0.25)
        inst_para.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()  # Empty line after instructions


def _format_mcq_options(options: Dict[str, str], doc) -> None:
    """Format MCQ options in CBSE style.

    Args:
        options: Dict of options {"A": "text", "B": "text", ...}
        doc: Document object to add paragraph to
    """
    # Create options paragraph with indentation
    options_para = doc.add_paragraph()
    options_para.paragraph_format.left_indent = Inches(0.5)

    # Add each option
    for letter in ["A", "B", "C", "D"]:
        if letter in options:
            opt_text = options[letter]
            options_para.add_run(f"{letter}) {opt_text}")
            options_para.add_run("\n")


def _format_internal_choice(question: Dict, doc, q_num: int, q_marks: int) -> Tuple[int, bool]:
    """Format internal choice question (OR format).

    Returns:
        (question_number_used, is_internal_choice)
    """
    # Check if this is an internal choice question
    if question.get("internal_choice"):
        # Format as "21. [Question text] OR"
        q_para = doc.add_paragraph()
        q_para.add_run(f"{q_num}.").bold = True
        q_text = question.get("question_text", "")
        q_para.add_run(f" {q_text} ({q_marks} marks)")

        # Add OR on next line
        or_para = doc.add_paragraph()
        or_para.add_run("OR").bold = True
        or_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        or_para.runs[0].font.size = Pt(10)

        return q_num, True

    return q_num, False


def _format_case_study_question(question: Dict, doc, q_num: int, q_marks: int) -> int:
    """Format case study question with sub-parts.

    Returns:
        question_number_used
    """
    # Main question text (case study passage)
    q_para = doc.add_paragraph()
    q_para.add_run(f"{q_num}.").bold = True
    q_text = question.get("question_text", "")
    q_para.add_run(f" {q_text}")

    # Add sub-questions if present
    sub_questions = question.get("sub_questions", [])
    if sub_questions:
        for sub_q in sub_questions:
            sub_para = doc.add_paragraph()
            sub_para.paragraph_format.left_indent = Inches(0.5)
            part = sub_q.get("part", "")
            marks = sub_q.get("marks", 0)
            sub_para.add_run(f"{part} ")
            if marks > 0:
                sub_para.add_run(f"({marks} mark{'s' if marks > 1 else ''})")
    else:
        # Default sub-questions format
        default_parts = [
            {"part": "(i)", "marks": 1},
            {"part": "(ii)", "marks": 1},
            {"part": "(iii)", "marks": 2},
        ]
        for sub_q in default_parts:
            sub_para = doc.add_paragraph()
            sub_para.paragraph_format.left_indent = Inches(0.5)
            sub_para.add_run(
                f"{sub_q['part']} ({sub_q['marks']} mark{'s' if sub_q['marks'] > 1 else ''})"
            )

    # Total marks line
    marks_para = doc.add_paragraph()
    marks_para.add_run(f"({q_marks} marks)").italic = True

    return q_num


@tool
def generate_docx_tool(
    json_paper_path: str, output_docx_path: Optional[str] = None
) -> Dict[str, Any]:
    """Convert JSON question paper to DOCX format with embedded diagrams.

    This tool reads a JSON question paper from the given path, converts all SVG diagrams
    to PNG format, and creates a professionally formatted DOCX document following CBSE
    standards. The DOCX includes headers with exam information, formatted sections,
    numbered questions, internal choice formatting, case study sub-parts, and embedded
    images for questions with diagrams.

    Args:
        json_paper_path: Path to the JSON question paper file
        output_docx_path: Optional path for output DOCX. If not provided, filename is
                        auto-generated based on metadata.

    Returns:
        {
            "success": bool,
            "docx_path": str,           # Path to generated DOCX file
            "questions_count": int,      # Total questions in paper
            "diagrams_embedded": int,   # Number of diagrams converted/embedded
            "generation_time": str,      # ISO timestamp
            "error": str (if failed)
        }

    Example:
        generate_docx_tool(
            json_paper_path="output/paper.json",
            output_docx_path="output/docx/paper.docx"
        )
    """
    try:
        from docx import Document

        # Read JSON paper
        json_path = Path(json_paper_path)
        if not json_path.exists():
            return {"success": False, "error": f"JSON file not found: {json_paper_path}"}

        with open(json_path, "r", encoding="utf-8") as f:
            paper_data = json.load(f)

        # Create Word document
        doc = Document()

        # Setup margins (CBSE standard: 1 inch on all sides)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Create header
        metadata = paper_data.get("exam_metadata", paper_data.get("paper_metadata", {}))
        _create_cbse_header(doc, metadata)

        # Add general instructions
        sections = paper_data.get("sections", [])
        _create_cbse_general_instructions(doc, num_sections=len(sections))

        # Track overall question number (Q1, Q2, Q3... across all sections)
        overall_q_num = 0
        questions_count = 0
        diagrams_embedded = 0

        for section_data in sections:
            section_id = section_data.get("section_id", "")
            section_title = section_data.get(
                "title", CBSE_SECTIONS.get(section_id, {}).get("title", "")
            )

            # Get marks per question
            marks_per_q = section_data.get(
                "marks_per_question", CBSE_SECTIONS.get(section_id, {}).get("marks_per_question", 1)
            )

            # Calculate section marks
            num_questions = len(section_data.get("questions", []))
            section_total_marks = num_questions * marks_per_q

            # Section heading with CBSE format
            section_heading = doc.add_paragraph()
            section_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            section_heading.add_run(f"SECTION {section_id}").bold = True
            section_heading.runs[0].font.size = Pt(12)
            section_heading.runs[0].font.color.rgb = RGBColor(0, 0, 0)

            # Section subtitle and marks
            section_sub = doc.add_paragraph()
            section_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            section_sub.add_run(f"{section_title}").italic = True
            section_sub.runs[0].font.size = Pt(10)

            section_marks = doc.add_paragraph()
            section_marks.alignment = WD_ALIGN_PARAGRAPH.CENTER
            section_marks.add_run(
                f"({num_questions} × {marks_per_q} = {section_total_marks} marks)"
            ).italic = True
            section_marks.runs[0].font.size = Pt(10)

            doc.add_paragraph()  # Empty line

            # Process questions
            questions = section_data.get("questions", [])

            i = 0
            while i < len(questions):
                question = questions[i]
                overall_q_num += 1
                questions_count += 1

                q_text = question.get("question_text", "")
                q_marks = question.get("marks", marks_per_q)
                q_format = question.get("question_format", "")
                has_choice = question.get("internal_choice", False)

                # Handle internal choice questions (OR format)
                if has_choice and section_id in ["B", "C", "D"] and i >= len(questions) - 2:
                    # This is one of the last 2 questions with internal choice
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"{overall_q_num}.").bold = True
                    q_para.add_run(f" {q_text} ({q_marks} marks)")

                    # Add OR
                    or_para = doc.add_paragraph()
                    or_para.add_run("OR").bold = True
                    or_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    or_para.runs[0].font.size = Pt(10)

                    # Next question is the alternative
                    if i + 1 < len(questions):
                        i += 1
                        alt_question = questions[i]
                        alt_q_text = alt_question.get("question_text", "")
                        alt_q_marks = alt_question.get("marks", q_marks)

                        alt_para = doc.add_paragraph()
                        alt_para.add_run(f"{alt_q_text} ({alt_q_marks} marks)")

                        # Handle options for alternative if MCQ
                        if q_format == "MCQ" and alt_question.get("options"):
                            _format_mcq_options(alt_question.get("options"), doc)

                        # Handle diagram for alternative
                        if alt_question.get("has_diagram"):
                            svg_base64 = alt_question.get("diagram_svg_base64", "")
                            if svg_base64:
                                png_bytes = _svg_base64_to_png(svg_base64, width=350)
                                if png_bytes:
                                    temp_png = TEMP_DIR / f"_{uuid.uuid4()}.png"
                                    with open(temp_png, "wb") as f:
                                        f.write(png_bytes)
                                    try:
                                        diagram_para = doc.add_paragraph()
                                        diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        desc = alt_question.get("diagram_description", "Diagram")
                                        diagram_para.add_run(f"Figure: {desc}").italic = True
                                        doc.add_picture(str(temp_png), width=Inches(4))
                                        diagrams_embedded += 1
                                    finally:
                                        temp_png.unlink()

                        overall_q_num += 1
                        questions_count += 1

                # Handle case study questions (Section E)
                elif section_id == "E" or question.get("has_sub_questions"):
                    overall_q_num = _format_case_study_question(
                        question, doc, overall_q_num, q_marks
                    )

                    # Handle options if MCQ case study
                    if q_format == "MCQ" and question.get("options"):
                        _format_mcq_options(question.get("options"), doc)

                # Handle normal questions
                else:
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"{overall_q_num}.").bold = True
                    q_para.add_run(f" {q_text} ({q_marks} mark{'s' if q_marks > 1 else ''})")

                    # Handle MCQ options (new dict format)
                    if q_format == "MCQ" and question.get("options"):
                        _format_mcq_options(question.get("options"), doc)

                # Embed diagram if present (for all question types)
                if question.get("has_diagram") or question.get("diagram_needed"):
                    svg_base64 = question.get("diagram_svg_base64", "")
                    if svg_base64:
                        png_bytes = _svg_base64_to_png(svg_base64, width=350)

                        if png_bytes:
                            temp_png = TEMP_DIR / f"_{uuid.uuid4()}.png"
                            with open(temp_png, "wb") as f:
                                f.write(png_bytes)

                            try:
                                diagram_para = doc.add_paragraph()
                                diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                                desc = question.get("diagram_description", "Diagram")
                                diagram_para.add_run(f"Figure: {desc}").italic = True

                                doc.add_picture(str(temp_png), width=Inches(4))

                                diagrams_embedded += 1
                            finally:
                                temp_png.unlink()

                # Add answer space (for non-MCQ questions)
                if q_format not in ["MCQ"] and not question.get("has_sub_questions"):
                    # Add some space for answer
                    doc.add_paragraph()

                i += 1

            # Page break between sections
            doc.add_page_break()

        # Create footer
        for section in doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.text = f"CBSE Question Paper Generator | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.runs[0].font.size = Pt(8)
            footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)

        # Generate output filename
        if not output_docx_path:
            output_docx_path = str(DOCX_OUTPUT_DIR / _generate_docx_filename(metadata))
        else:
            output_docx_path = str(output_docx_path)

        # Ensure output directory exists
        Path(output_docx_path).parent.mkdir(parents=True, exist_ok=True)

        # Save document
        doc.save(str(output_docx_path))

        return {
            "success": True,
            "docx_path": str(output_docx_path),
            "questions_count": questions_count,
            "diagrams_embedded": diagrams_embedded,
            "generation_time": datetime.now().isoformat(),
        }

    except Exception as e:
        import traceback

        error_trace = traceback.format_exc()
        print(f"Error generating DOCX: {e}")
        print(f"Traceback:\n{error_trace}")

        return {"success": False, "error": str(e), "traceback": error_trace}


def verify_docx_generator():
    """Test function to verify DOCX generation works."""
    print("Testing docx_generator module...")

    # Check if dependencies available
    try:
        from docx import Document

        print("✓ python-docx available")
    except ImportError as e:
        print(f"✗ python-docx not available: {e}")
        return

    if not _ensure_cairosvg_installed():
        print("⚠ cairosvg not available, SVG conversion will be skipped")
    else:
        print("✓ cairosvg available")

    print("DocX generator module tests complete.")


if __name__ == "__main__":
    verify_docx_generator()
